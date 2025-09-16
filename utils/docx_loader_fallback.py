"""Fallback loader for .docx files using python-docx directly"""

import os
from typing import List
from docx import Document as DocxDocument
from langchain.schema import Document

class SimpleDocxLoader:
    """Simple .docx loader using python-docx directly to avoid unstructured issues"""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        """Load .docx file and return as LangChain Document"""
        try:
            doc = DocxDocument(self.file_path)

            # Extract all text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(' | '.join(row_text))

            # Join all text
            content = '\n'.join(full_text)

            # Create metadata
            metadata = {
                "source": self.file_path,
                "file_name": os.path.basename(self.file_path),
                "file_type": "docx"
            }

            # Return as LangChain Document
            return [Document(page_content=content, metadata=metadata)]

        except Exception as e:
            print(f"Error loading {self.file_path}: {e}")
            return []